import streamlit as st
import os
import re
import json
import logging
from io import BytesIO
from datetime import datetime
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from typing import List, Dict, Any, Optional

# Try to import google.genai (new), fall back to google.generativeai (deprecated)
GENAI_CLIENT = None

def get_genai_client(api_key: str):
    """Get or create the GenAI client. Supports both old and new SDK."""
    global GENAI_CLIENT
    if GENAI_CLIENT is not None:
        return GENAI_CLIENT
    try:
        # Try new SDK first
        from google import genai
        GENAI_CLIENT = genai.Client(api_key=api_key)
        return GENAI_CLIENT
    except (ImportError, AttributeError):
        pass
    try:
        # Fall back to old SDK
        from google import generativeai as genai
        genai.configure(api_key=api_key)
        GENAI_CLIENT = genai
        return GENAI_CLIENT
    except (ImportError, AttributeError):
        return None

# Configure logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

# --- HẰNG SỐ (CONSTANTS) ---

SYSTEM_INSTRUCTION = (
    "Bạn là một nhà phân tích tài liệu kỹ thuật. Nhiệm vụ của bạn là trích xuất thông tin từ 'Biên bản bàn giao' "
    "vào định dạng JSON. "
    "QUAN TRỌNG: Trường 'pk' (Phụ kiện) phải là một danh sách (Array) các chuỗi, không được gộp thành 1 chuỗi dài. "
    "Nếu không có thông tin, trả về null. Không thêm Markdown (```json)."
)

CONFIG_FILE_PATH = 'config.ini'
TEMPLATE_FILE = 'bbbg.docx'

DESIRED_MODELS_KEYWORDS = ['gemini', 'gemma', 'flash', 'pro']
EXCLUDE_MODELS_KEYWORDS = ['bison', 'gecko', 'embedding', 'aqa', 'vision', 'legacy']

MAX_FILENAME_LEN = 200
MAX_SERI_DISPLAY = 100
MAX_DEVICES_IN_FILENAME = 2
DEFAULT_FONT_NAME = 'Times New Roman'
DEFAULT_FONT_SIZE = 12

# API Keys for rotation
API_KEYS = [
    "AIzaSyBlg4BpQn9UdydetlS3ycHdkX4i231k7Yg",
]
_current_key_idx = 0

# --- HELPER FUNCTIONS ---

def convert_none_to_empty_string(obj: Any) -> Any:
    if isinstance(obj, dict):
        return {k: convert_none_to_empty_string(v) for k, v in obj.items()}
    if isinstance(obj, list):
        return [convert_none_to_empty_string(elem) for elem in obj]
    return "" if obj is None else obj

def clean_filename(filename: str) -> str:
    chars_to_remove = r'[\\/*?":<>|.]'
    cleaned = re.sub(chars_to_remove, '', filename)
    return cleaned[:MAX_FILENAME_LEN] if len(cleaned) > MAX_FILENAME_LEN else cleaned

def standardize_string(text: Any) -> str:
    if not isinstance(text, str):
        return str(text)

    text = text.replace('Ằ', 'Ă').replace('Ắ', 'Ă').replace('Ặ', 'Ă').replace('Ẳ', 'Ă').replace('Ẵ', 'Ă')
    text = text.replace('È', 'E').replace('É', 'E').replace('Ẹ', 'E').replace('Ẻ', 'E').replace('Ẽ', 'E')
    text = text.replace('Ề', 'E').replace('Ế', 'E').replace('Ệ', 'E').replace('Ể', 'E').replace('Ễ', 'E')
    text = text.replace('Ì', 'I').replace('Í', 'I').replace('Ị', 'I').replace('Ỉ', 'I').replace('Ĩ', 'I')
    text = text.replace('Ò', 'O').replace('Ó', 'O').replace('Ọ', 'O').replace('Ỏ', 'O').replace('Õ', 'O')
    text = text.replace('Ồ', 'O').replace('Ố', 'O').replace('Ộ', 'O').replace('Ổ', 'O').replace('Ỗ', 'O')
    text = text.replace('Ờ', 'O').replace('Ớ', 'O').replace('Ợ', 'O').replace('Ở', 'O').replace('Ỡ', 'O')
    text = text.replace('Ù', 'U').replace('Ú', 'U').replace('Ụ', 'U').replace('Ủ', 'U').replace('Ũ', 'U')
    text = text.replace('Ừ', 'U').replace('Ứ', 'U').replace('Ự', 'U').replace('Ử', 'U').replace('Ữ', 'U')
    text = text.replace('Ỳ', 'Y').replace('Ý', 'Y').replace('Ỵ', 'Y').replace('Ỷ', 'Y').replace('Ỹ', 'Y')
    text = text.replace('Đ', 'D')
    text = text.lower().replace('-', ' ').strip()
    return re.sub(r'\s+', ' ', text).strip()

def shorten_company_name(company_name: str) -> str:
    if not isinstance(company_name, str):
        return str(company_name).strip()

    original = company_name.strip()
    name = original

    prefixes = [
        r"CÔNG TY TNHH MỘT THÀNH VIÊN", r"CÔNG TY TNHH MTV", r"CÔNG TY TNHH HAI THÀNH VIÊN TRỞ LÊN",
        r"CÔNG TY CỔ PHẦN", r"CÔNG TY TNHH", r"CÔNG TY", r"TNHH", r"CỔ PHẦN",
    ]
    suffixes = [
        r"MỘT THÀNH VIÊN", r"MTV", r"HAI THÀNH VIÊN TRỞ LÊN", r"CỔ PHẦN", r"TNHH",
    ]
    common_terms = [
        r"THƯƠNG MẠI VÀ DỊCH VỤ", r"DỊCH VỤ VÀ THƯƠNG MẠI", r"TM VÀ DV", r"DV VÀ TM", r"TM & DV", r"DV & TM",
        r"TM", r"DV", r"CÔNG NGHỆ", r"THƯƠNG MẠI", r"TRANG THIẾT BỊ", r"Y TẾ", r"XÂY DỰNG",
        r"ĐẦU TƯ", r"PHÁT TRIỂN", r"GIẢI PHÁP", r"KỸ THUẬT", r"SẢN XUẤT", r"NHẬP KHẨU", r"XUẤT NHẬP KHẨU",
        r"KINH DOANH", r"PHÂN PHỐI", r"VIỆT NAM"
    ]

    for p in prefixes + suffixes:
        name = re.sub(r'^\s*' + re.escape(p) + r'\s*|\s*' + re.escape(p) + r'\s*$', '', name, flags=re.IGNORECASE).strip(" ,.-_&")

    for term in common_terms:
        name = re.sub(r'\b' + re.escape(term) + r'\b', '', name, flags=re.IGNORECASE).strip()
        name = re.sub(r'\s+', ' ', name).strip(" ,.-_&")

    return name if name else original

# --- CORE LOGIC ---

def group_devices(device_list: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    grouped: Dict[tuple, Dict[str, Any]] = {}

    for item in device_list:
        if not isinstance(item, dict):
            continue

        raw_pk = item.get('pk', '')
        pk_key = json.dumps(raw_pk, ensure_ascii=False, sort_keys=True) if isinstance(raw_pk, list) else str(raw_pk).strip() if raw_pk else ''

        group_key = (
            standardize_string(item.get('ttb', '')).strip(),
            str(item.get('model', '')).strip(),
            str(item.get('hang', '')).strip(),
            str(item.get('nsx', '')).strip(),
            str(item.get('dvt', '')).strip(),
            pk_key
        )

        try:
            sl_raw = item.get('sl', '0')
            cleaned_sl = re.sub(r'[^\d.]', '', str(sl_raw).strip())
            current_sl = float(cleaned_sl) if cleaned_sl else 0
        except (ValueError, TypeError):
            current_sl = 0

        seri_raw = item.get('seri', [])
        if isinstance(seri_raw, str):
            seri_list = [seri_raw] if seri_raw else []
        elif isinstance(seri_raw, list):
            seri_list = [str(s).strip() for s in seri_raw if s and str(s).strip()]
        else:
            seri_list = [str(seri_raw)] if seri_raw else []

        if group_key not in grouped:
            grouped[group_key] = {
                'ttb': str(item.get('ttb', '')).strip(),
                'model': str(item.get('model', '')).strip(),
                'hang': str(item.get('hang', '')).strip(),
                'nsx': str(item.get('nsx', '')).strip(),
                'dvt': str(item.get('dvt', '')).strip(),
                'pk_raw': raw_pk,
                'total_sl': current_sl,
                'seri': set(seri_list)
            }
        else:
            grouped[group_key]['total_sl'] += current_sl
            grouped[group_key]['seri'].update(seri_list)

    final_list = []
    for group_data in grouped.values():
        unique_seri = sorted(group_data['seri']) if group_data['seri'] else []
        display_seri = unique_seri[:MAX_SERI_DISPLAY]
        seri_text = f"Số seri: {', '.join(display_seri)}"
        if len(unique_seri) > MAX_SERI_DISPLAY:
            seri_text += f" (và {len(unique_seri) - MAX_SERI_DISPLAY} seri khác)"

        final_list.append({
            'ttb': group_data['ttb'],
            'model': group_data['model'],
            'hang': group_data['hang'],
            'nsx': group_data['nsx'],
            'dvt': group_data['dvt'],
            'sl': group_data['total_sl'],
            'pk': group_data['pk_raw'],
            'seri_text': seri_text
        })

    return final_list

def generate_filename(data: Dict[str, Any], grouped_devices: List[Dict[str, Any]]) -> str:
    device_parts = []
    for item in grouped_devices[:MAX_DEVICES_IN_FILENAME]:
        quantity = int(item.get('sl', 0))
        formatted_quantity = f"{quantity:02d}"
        device_name = str(item.get('ttb', '')).strip()
        cleaned_device_name = re.sub(r'[\\/*?":<>|{}\[\]().,_]', '', device_name).strip()
        if cleaned_device_name:
            device_parts.append(f"{formatted_quantity} {cleaned_device_name}")
    device_info_str = "-".join(device_parts) or "ThietBi"

    cty_name_full = str(data.get('cty', 'UnknownCompany')).strip()
    cleaned_cty_name = shorten_company_name(cty_name_full)
    if not cleaned_cty_name:
        cleaned_cty_name = re.sub(r'[\\/*?":<>|{}\[\]()]', '', cty_name_full).strip(" ,.-_&") or "CongTy"

    shd_value = str(data.get('shd', '')).strip()
    shd_main_part = shd_value.split('-', 1)[0].strip() or "SoDinhDanh"
    shd_cleaned = clean_filename(shd_main_part)

    raw_filename = f"{device_info_str}_{cleaned_cty_name}_{shd_cleaned}"
    final_filename_base = re.sub(r'\s+', '_', clean_filename(raw_filename)).strip('_')

    if not final_filename_base or len(final_filename_base) < 3:
        return f"BienBanBanGiao_{cleaned_cty_name}_{shd_cleaned}.docx"

    return final_filename_base + '.docx'

def format_accessories_list(pk_raw: Any) -> str:
    if not pk_raw:
        return ""

    pk_lines = []
    if isinstance(pk_raw, list):
        pk_lines = [str(x).strip() for x in pk_raw if x]
    elif isinstance(pk_raw, str) and pk_raw:
        clean_str = re.sub(r'(cấu hình bao gồm|bao gồm|chi tiết cấu hình):', '', pk_raw, flags=re.IGNORECASE).replace('–', '-').strip()
        pk_lines = re.split(r'[;\n]+', clean_str)

    formatted = []
    for acc in pk_lines:
        clean_acc = acc.strip().lstrip('-•+').strip()
        if clean_acc:
            formatted.append(f"  + {clean_acc}")

    return "\n- Phụ kiện:\n" + "\n".join(formatted) if formatted else ""

def fill_word_template(data: Dict[str, Any], grouped_devices: List[Dict[str, Any]]) -> BytesIO:
    try:
        document = Document(TEMPLATE_FILE)
    except Exception as e:
        logger.error(f"Failed to open template: {e}")
        st.error(f"❌ Lỗi mở file mẫu '{TEMPLATE_FILE}': {e}", icon="❌")
        raise

    try:
        table = document.tables[0]
        for i in range(len(table.rows) - 1, 0, -1):
            table.rows[i]._element.getparent().remove(table.rows[i]._element)

        for count, item in enumerate(grouped_devices, 1):
            pk_text = format_accessories_list(item.get('pk'))
            device_info = (
                f"{item.get('ttb', '').strip()}\n"
                f"- Model: {item.get('model', '').strip()}\n"
                f"- Hãng: {item.get('hang', '').strip()}\n"
                f"- NSX: {item.get('nsx', '').strip()}"
                f"{pk_text}"
            )

            new_row = table.add_row()
            row_data = [
                str(count),
                device_info,
                item.get('dvt', '').strip(),
                str(int(item.get('sl', 0))),
                item.get('seri_text', '')
            ]

            for i, cell_text in enumerate(row_data):
                cell = new_row.cells[i]
                cell.text = str(cell_text)
                alignment = WD_ALIGN_PARAGRAPH.CENTER if i in (0, 2, 3) else WD_ALIGN_PARAGRAPH.LEFT
                for para in cell.paragraphs:
                    para.alignment = alignment
                    for run in para.runs:
                        run.font.name = DEFAULT_FONT_NAME
                        run.font.size = Pt(DEFAULT_FONT_SIZE)

    except IndexError:
        st.error("❌ File mẫu không có bảng.", icon="❌")
        raise

    now = datetime.now()
    shd_value = str(data.get('shd', '')).strip()
    shd_type = str(data.get('shd_type', 'Khác')).strip()

    shd_replacement = ""
    if shd_value:
        shd_type_lower = standardize_string(shd_type)
        if 'hop dong' in shd_type_lower or 'hd' in shd_type_lower:
            shd_replacement = f"Dựa theo HĐ số: {shd_value}"
        elif 'po' in shd_type_lower or 'de nghi' in shd_type_lower:
            shd_replacement = f"Dựa theo PO: {shd_value}"
        else:
            shd_replacement = f"Dựa theo số: {shd_value}"

    shd_pattern = re.compile(re.escape("shd"), re.IGNORECASE)

    for para in document.paragraphs:
        for run in para.runs:
            if "day" in run.text:
                run.text = run.text.replace("day", str(now.day))
            if "month" in run.text:
                run.text = run.text.replace("month", str(now.month))
            if "year" in run.text:
                run.text = run.text.replace("year", str(now.year))
            if shd_pattern.search(run.text):
                run.text = shd_pattern.sub(shd_replacement, run.text)

    byte_io = BytesIO()
    document.save(byte_io)
    byte_io.seek(0)
    return byte_io

# --- API CONFIG ---

def get_api_key() -> Optional[str]:
    """Get API key from environment variable (Streamlit Cloud) or config file."""
    global _current_key_idx

    # Collect all available keys
    keys = []

    # Try Streamlit secrets first (for cloud deployment)
    try:
        import streamlit as st
        if hasattr(st, 'secrets'):
            for key in ['GEMINI_API_KEY', 'GEMINI_API_KEY_2', 'GEMINI_API_KEY_3',
                       'GEMINI_API_KEY_4', 'GEMINI_API_KEY_5', 'GEMINI_API_KEY_6',
                       'GEMINI_API_KEY_7', 'GEMINI_API_KEY_8']:
                if key in st.secrets:
                    val = st.secrets[key]
                    if val and val not in keys:
                        keys.append(val)
    except Exception:
        pass

    # Environment variable fallback
    api_key = os.environ.get('GEMINI_API_KEY')
    if api_key and api_key not in keys:
        keys.append(api_key)
    for i in range(2, 9):
        backup_key = os.environ.get(f'GEMINI_API_KEY_{i}')
        if backup_key and backup_key not in keys:
            keys.append(backup_key)

    # Config file (local development)
    if os.path.exists(CONFIG_FILE_PATH):
        try:
            import configparser
            config = configparser.ConfigParser()
            config.read(CONFIG_FILE_PATH)
            key = config['API']['GEMINI_API_KEY']
            if key and key != 'YOUR_API_KEY_HERE' and key not in keys:
                keys.append(key)
        except Exception:
            pass

    # Use indexed key from rotation pool
    if keys:
        return keys[_current_key_idx % len(keys)]

    # Fallback to hardcoded keys
    if _current_key_idx < len(API_KEYS):
        return API_KEYS[_current_key_idx]
    return None

def rotate_api_key() -> Optional[str]:
    global _current_key_idx

    # Collect all available keys
    keys = []

    # Try Streamlit secrets first (for cloud deployment)
    try:
        import streamlit as st
        if hasattr(st, 'secrets'):
            for key in ['GEMINI_API_KEY', 'GEMINI_API_KEY_2', 'GEMINI_API_KEY_3',
                       'GEMINI_API_KEY_4', 'GEMINI_API_KEY_5', 'GEMINI_API_KEY_6',
                       'GEMINI_API_KEY_7', 'GEMINI_API_KEY_8']:
                if key in st.secrets:
                    val = st.secrets[key]
                    if val and val not in keys:
                        keys.append(val)
    except Exception:
        pass

    # Environment variable fallback
    api_key = os.environ.get('GEMINI_API_KEY')
    if api_key and api_key not in keys:
        keys.append(api_key)
    for i in range(2, 9):
        backup_key = os.environ.get(f'GEMINI_API_KEY_{i}')
        if backup_key and backup_key not in keys:
            keys.append(backup_key)

    # Config file fallback
    if os.path.exists(CONFIG_FILE_PATH):
        try:
            import configparser
            config = configparser.ConfigParser()
            config.read(CONFIG_FILE_PATH)
            key = config['API']['GEMINI_API_KEY']
            if key and key != 'YOUR_API_KEY_HERE' and key not in keys:
                keys.append(key)
        except Exception:
            pass

    if not keys:
        keys = API_KEYS

    _current_key_idx = (_current_key_idx + 1) % len(keys)
    return keys[_current_key_idx] if keys else None

def configure_api(key: str) -> bool:
    """Configure the GenAI client with the given API key."""
    global GENAI_CLIENT
    GENAI_CLIENT = None  # Reset to force re-initialization
    try:
        client = get_genai_client(key)
        return client is not None
    except Exception as e:
        logger.warning(f"Failed to configure API key: {e}")
        return False

def list_models() -> List[str]:
    try:
        client = get_genai_client(get_api_key() or API_KEYS[0])
        if client is None:
            return []
        # New SDK: client.models.list()
        if hasattr(client, 'models'):
            return [m.name for m in client.models.list() if 'generateContent' in getattr(m, 'supported_generation_methods', [])]
        # Old SDK: genai.list_models()
        return [m.name for m in client.list_models() if 'generateContent' in m.supported_generation_methods]
    except Exception:
        return []

@st.cache_resource
def check_prerequisites() -> bool:
    """Kiểm tra API key và file template."""
    api_key = get_api_key()
    if not api_key:
        st.error("❌ Không tìm thấy GEMINI_API_KEY.", icon="❌")
        return False

    if not configure_api(api_key):
        # Try all available keys
        for _ in range(8):
            next_key = rotate_api_key()
            if next_key and configure_api(next_key):
                break
        else:
            st.error("❌ Không thể cấu hình API key.", icon="❌")
            return False

    try:
        models = list_models()
        if not models:
            logger.warning("Could not list models, using fallback model list")
    except Exception as e:
        logger.warning(f"Could not list models: {e}")

    if not os.path.exists(TEMPLATE_FILE):
        st.error(f"❌ Thiếu file mẫu '{TEMPLATE_FILE}'", icon="❌")
        return False

    return True

@st.cache_data
def get_filtered_models() -> List[str]:
    logger.info("Getting filtered models...")
    models = list_models()
    logger.info(f"list_models returned {len(models)} models")
    found = []
    for m in models:
        name = m.lower()
        if any(k in name for k in DESIRED_MODELS_KEYWORDS) and not any(k in name for k in EXCLUDE_MODELS_KEYWORDS):
            found.append(m)

    # Fallback to known working models if none found
    if not found:
        logger.warning("No models found from API, using fallback list")
        found = [
            'gemini-3-flash',
            'gemini-3.1-flash-lite',
            'gemini-2.5-flash',
            'gemini-2.5-flash-lite',
            'gemini-2.0-flash',
            'gemma-4-31b-it',
            'gemma-4-26b-it',
            'gemini-2.5-pro',
        ]

    def priority(nm):
        n = nm.lower()
        # Higher RPD models first (gemini-3.1-flash-lite has 500 RPD)
        if "gemini-3.1-flash-lite" in n: return 0
        if "gemini-2.5-flash-lite" in n: return 1
        if "gemini-3-flash" in n: return 2
        if "gemini-2.5-flash" in n: return 3
        if "gemini-2.0-flash" in n: return 4
        if "gemma-4-31b" in n: return 5
        if "gemma-4-26b" in n: return 6
        if "gemini-2.5-pro" in n: return 7
        return 8

    found.sort(key=priority)
    logger.info(f"Returning models: {found}")
    return found

def call_gemini_vision_api(uploaded_file_bytes: bytes, mime_type: str, prompt: str, model_list: List[str]) -> Optional[Dict]:
    if not model_list:
        st.error("Không có model nào khả dụng.", icon="❌")
        return None

    global _current_key_idx
    last_error = None

    # Get actual count of available keys
    available_keys = []
    try:
        import streamlit as st
        if hasattr(st, 'secrets'):
            for key in ['GEMINI_API_KEY', 'GEMINI_API_KEY_2', 'GEMINI_API_KEY_3',
                       'GEMINI_API_KEY_4', 'GEMINI_API_KEY_5', 'GEMINI_API_KEY_6',
                       'GEMINI_API_KEY_7', 'GEMINI_API_KEY_8']:
                if key in st.secrets:
                    val = st.secrets[key]
                    if val and val not in available_keys:
                        available_keys.append(val)
    except Exception:
        pass

    env_key = os.environ.get('GEMINI_API_KEY')
    if env_key and env_key not in available_keys:
        available_keys.append(env_key)
    for i in range(2, 9):
        backup_key = os.environ.get(f'GEMINI_API_KEY_{i}')
        if backup_key and backup_key not in available_keys:
            available_keys.append(backup_key)

    if os.path.exists(CONFIG_FILE_PATH):
        try:
            import configparser
            config = configparser.ConfigParser()
            config.read(CONFIG_FILE_PATH)
            key = config['API']['GEMINI_API_KEY']
            if key and key != 'YOUR_API_KEY_HERE' and key not in available_keys:
                available_keys.append(key)
        except Exception:
            pass

    if not available_keys:
        available_keys = API_KEYS

    key_count = len(available_keys)

    # Collect all unique models across all keys for comprehensive retry
    all_models = []
    for key in available_keys:
        all_models.extend(model_list)
    max_attempts = len(all_models) + key_count  # Model attempts + key rotations

    attempt = 0
    while attempt < max_attempts:
        api_key = get_api_key()
        if not api_key:
            break

        if not configure_api(api_key):
            rotate_api_key()
            attempt += 1
            continue

        for model_name in model_list:
            try:
                with st.spinner(f"✨ Đang dùng model: {model_name}..."):
# Try new SDK first
                    try:
                        client = get_genai_client(get_api_key() or available_keys[0])
                        if client is None:
                            continue

                        # New SDK path
                        from google.genai import types as genai_types
                        from google.genai.types import PartDict
                        
                        file_part = PartDict(mime_type=mime_type, binary=uploaded_file_bytes)
                        
                        # Build contents - use positional arg for model name
                        contents = [
                            {
                                "role": "user",
                                "parts": [
                                    {"text": prompt},
                                    file_part
                                ]
                            }
                        ]
                        
                        # New SDK uses model= as positional, not model_name=
                        response = client.models.generate_content(
                            model=model_name,
                            contents=contents,
                            config=genai_types.GenerateContentConfig(
                                system_instruction=SYSTEM_INSTRUCTION
                            )
                        )
                        text = response.text.strip()
                    except (ImportError, AttributeError):
                        # Fall back to old SDK
                        from google.generativeai import types as genai_types_old
                        from google.generativeai import GenerativeModel
                        
                        file_part = genai_types_old.Blob(mime_type=mime_type, data=uploaded_file_bytes)
                        model = GenerativeModel(model_name=model_name, system_instruction=SYSTEM_INSTRUCTION)
                        response = model.generate_content([file_part, prompt])
                        text = response.text.strip()

                    if text.startswith("```json"):
                        text = text[7:]
                    elif text.startswith("```"):
                        text = text[3:]
                    if text.endswith("```"):
                        text = text[:-3]

                    data = json.loads(text.strip())
                    st.success(f"✅ Thành công với model: {model_name}")
                    logger.info(f"Successfully extracted data using model: {model_name}")
                    return data

            except json.JSONDecodeError as e:
                logger.warning(f"Model {model_name} returned non-JSON: {e}")
                last_error = f"Model {model_name} trả về dữ liệu không hợp lệ"
                continue

            except Exception as e:
                err_str = str(e)
                err_type = type(e).__name__
                logger.warning(f"Model {model_name} failed ({err_type}): {err_str}")

                # Check for quota/key errors
                quota_errors = ["API_KEY", "UNAUTHORIZED", "INVALID", "quota", "limit", "429",
                               "RESOURCE_EXHAUSTED", "ResourceExhausted", "leaked", "PERMISSION_DENIED"]
                if any(x in err_str.upper() for x in quota_errors) or "429" in err_str:
                    logger.warning(f"Quota/key error detected, rotating to next key...")
                    rotate_api_key()
                    attempt += 1
                    # Brief pause to avoid hammering
                    import time
                    time.sleep(0.5)
                    break

                last_error = err_str
                continue

    if last_error:
        st.error(f"❌ Tất cả model và key đều thất bại: {last_error}", icon="❌")
    else:
        st.error("❌ Không thể kết nối API. Vui lòng kiểm tra API key và thử lại.", icon="❌")
    return None

# --- MAIN ---

PROMPT_TEMPLATE = """
Hãy trích xuất thông tin từ Biên bản bàn giao và trả về JSON hợp lệ.

**Cấu trúc JSON bắt buộc:**
{
  "shd": "Số định danh (số hợp đồng, PO, v.v.)",
  "shd_type": "Loại: 'Hợp đồng', 'PO', 'Đề nghị', hoặc 'Khác'",
  "cty": "Tên công ty bên giao",
  "ds": [
    {
      "ttb": "Tên thiết bị",
      "model": "Model",
      "hang": "Hãng",
      "nsx": "Nước sản xuất",
      "dvt": "Đơn vị tính",
      "sl": Số lượng (số nguyên),
      "seri": ["danh sách số seri"] hoặc null,
      "pk": ["danh sách phụ kiện"] hoặc null
    }
  ]
}

**Quy tắc quan trọng:**
1. pk PHẢI là một ARRAY các chuỗi, ví dụ: ["Dây nguồn", "Cáp USB"]
2. pk KHÔNG được gộp thành một chuỗi dài
3. Nếu không có thông tin, trả về null
4. KHÔNG có Markdown code block, chỉ trả về JSON thuần
5. Đọc CHÍNH XÁC model/model number - cẩn thận với các số giống nhau (VD: 0/O, 1/I/l, 2/Z, 5/S, 6/G, 8/B)
6. Nếu không chắc chắn, ghi lại như trong tài liệu
"""

def main():
    st.set_page_config(
        page_title="Chuyển đổi Bàn giao",
        page_icon="📋",
        layout="centered"
    )

    st.markdown("""
    <style>
    .stFileUploader { border: 2px dashed #004aad; padding: 1rem; border-radius: 0.5rem; }
    .stDownloadButton > button { background-color: #4CAF50; color: white; font-weight: bold; }
    h1 { color: #004aad; text-align: center; }
    </style>
    """, unsafe_allow_html=True)

    st.title("📋 Chuyển đổi Biên bản Bàn giao")
    st.markdown("Tải lên file PDF hoặc ảnh từ Biên bản bàn giao để tạo file Word")

    if not check_prerequisites():
        st.stop()

    available_models = get_filtered_models()
    if not available_models:
        st.error("Không tìm thấy model Gemini phù hợp.", icon="❌")
        st.stop()

    st.success(f"✅ Đã kết nối thành công! Tìm thấy {len(available_models)} model.")

    uploaded_file = st.file_uploader(
        "📁 Chọn file (PDF/PNG/JPG)",
        type=["pdf", "jpg", "png"],
        help="Hỗ trợ định dạng PDF, PNG, JPG"
    )

    if uploaded_file:
        st.info(f"📥 Đang xử lý: **{uploaded_file.name}**", icon="⏳")

        file_bytes = uploaded_file.getvalue()
        mime = 'application/pdf' if uploaded_file.name.lower().endswith('.pdf') else 'image/jpeg'

        data = call_gemini_vision_api(
            file_bytes,
            mime,
            PROMPT_TEMPLATE,
            available_models
        )

        if data and 'ds' in data:
            data = convert_none_to_empty_string(data)
            grouped = group_devices(data['ds'])

            filename = generate_filename(data, grouped)
            word_io = fill_word_template(data, grouped)

            st.download_button(
                "⬇️ Tải xuống file Word",
                word_io,
                filename,
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
            st.success("✅ Hoàn thành! File Word đã sẵn sàng tải xuống.", icon="🎉")
        else:
            st.error("❌ Không trích xuất được dữ liệu từ file. Vui lòng thử file khác.", icon="❌")

if __name__ == "__main__":
    main()