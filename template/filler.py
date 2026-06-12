"""Word template filling — populates bbbg.docx with extracted data."""

import re
from io import BytesIO
from datetime import datetime
from typing import Dict, Any, List
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from core.models import GroupedDevice
from utils.logging_setup import get_logger

logger = get_logger('template.filler')

TEMPLATE_FILE = 'bbbg.docx'
DEFAULT_FONT_NAME = 'Times New Roman'
DEFAULT_FONT_SIZE = 12


def format_accessories_list(pk_raw: Any) -> str:
    """Format accessories (pk) into a bullet list string for Word cell."""
    if not pk_raw:
        return ""

    pk_lines = []
    if isinstance(pk_raw, list):
        pk_lines = [str(x).strip() for x in pk_raw if x]
    elif isinstance(pk_raw, str) and pk_raw:
        clean_str = re.sub(
            r'(cấu hình bao gồm|bao gồm|chi tiết cấu hình):',
            '', pk_raw, flags=re.IGNORECASE
        ).replace('–', '-').strip()
        pk_lines = re.split(r'[;\n]+', clean_str)

    formatted = []
    for acc in pk_lines:
        clean_acc = acc.strip().lstrip('-•+').strip()
        if clean_acc:
            formatted.append(f"  + {clean_acc}")

    return "\n- Phụ kiện:\n" + "\n".join(formatted) if formatted else ""


def fill_word_template(
    data: Dict[str, Any],
    grouped_devices: List[GroupedDevice],
) -> BytesIO:
    """Fill the Word template with handover data and grouped devices."""
    try:
        document = Document(TEMPLATE_FILE)
    except Exception as e:
        logger.error(f"Failed to open template: {e}")
        raise

    try:
        table = document.tables[0]
        for i in range(len(table.rows) - 1, 0, -1):
            table.rows[i]._element.getparent().remove(table.rows[i]._element)

        for count, item in enumerate(grouped_devices, 1):
            pk_text = format_accessories_list(item.pk)
            device_info = (
                f"{item.ttb.strip()}\n"
                f"- Model: {item.model.strip()}\n"
                f"- Hãng: {item.hang.strip()}\n"
                f"- NSX: {item.nsx.strip()}"
                f"{pk_text}"
            )

            new_row = table.add_row()
            row_data = [
                str(count),
                device_info,
                item.dvt.strip(),
                str(int(item.sl)),
                item.seri_text,
            ]

            for i, cell_text in enumerate(row_data):
                cell = new_row.cells[i]
                cell.text = str(cell_text)
                alignment = WD_ALIGN_PARAGRAPH.CENTER if i in (0, 2, 3) else WD_ALIGN_PARAGRAPH.LEFT
                for para in cell.paragraphs:
                    para.alignment = alignment
                    for run in para.runs:
                        run.font.name = DEFAULT_FONT_NAME
                        run.font.size = Pt(DEFAULT_FONT_SIZE)

    except IndexError:
        logger.error("Template file has no table")
        raise

    now = datetime.now()
    shd_value = str(data.get('shd', '')).strip()
    shd_type = str(data.get('shd_type', 'Khác')).strip()

    shd_replacement = ""
    if shd_value:
        shd_type_lower = shd_type.lower().replace(' ', '')
        if 'hopdong' in shd_type_lower or 'hd' in shd_type_lower:
            shd_replacement = f"Dựa theo HĐ số: {shd_value}"
        elif 'po' in shd_type_lower or 'de nghi' in shd_type_lower:
            shd_replacement = f"Dựa theo PO: {shd_value}"
        else:
            shd_replacement = f"Dựa theo số: {shd_value}"

    shd_pattern = re.compile(re.escape("shd"), re.IGNORECASE)

    for para in document.paragraphs:
        for run in para.runs:
            if "day" in run.text:
                run.text = run.text.replace("day", str(now.day))
            if "month" in run.text:
                run.text = run.text.replace("month", str(now.month))
            if "year" in run.text:
                run.text = run.text.replace("year", str(now.year))
            if shd_pattern.search(run.text):
                run.text = shd_pattern.sub(shd_replacement, run.text)

    byte_io = BytesIO()
    document.save(byte_io)
    byte_io.seek(0)
    return byte_io
