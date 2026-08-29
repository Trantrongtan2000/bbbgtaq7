import pytest
import docx
from core.models import GroupedDevice
from template.filler import fill_word_template, format_quantity, format_accessories_list

def test_format_quantity():
 assert format_quantity(1.0) == '1'
 assert format_quantity(1.5) == '1.5'
 assert format_quantity(2.50) == '2.5'
 assert format_quantity(0) == '0'
 assert format_quantity('3') == '3'

def test_format_accessories_list():
    pk_list = ['Day nguon', 'Cap sensor']
    formatted = format_accessories_list(pk_list)
    assert 'Phụ kiện' in formatted
    assert '+ Day nguon' in formatted
    assert '+ Cap sensor' in formatted

def test_fill_word_template_execution():
 data = {
 'shd': '12345/HD',
 'shd_type': 'Hop dong',
 'cty': 'Cong ty Test'
 }
 devices = [
 GroupedDevice(
 ttb='May theo doi benh nhan',
 model='ePM 10',
 ref='REF-99',
 hang='Mindray',
 nsx='Trung Quoc',
 dvt='May',
 sl=1.5,
 pk=['Cap ECG'],
 seri_text='So seri: SN001'
 )
 ]
 bio = fill_word_template(data, devices)
 assert bio.getvalue() is not None
 doc = docx.Document(bio)
 assert len(doc.tables[0].rows) == 2
 row_text = ' '.join(c.text for c in doc.tables[0].rows[1].cells)
 assert '1.5' in row_text
 assert 'ePM 10' in row_text
 assert 'REF-99' in row_text
